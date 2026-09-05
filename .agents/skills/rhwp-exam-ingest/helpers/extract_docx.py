#!/usr/bin/env python3
"""rhwp-exam-ingest helper: DOCX → 텍스트 + 임베디드 이미지 추출.

사용법:
    extract_docx.py <input.docx> <out_dir>
    extract_docx.py --dry-run <input.docx> <out_dir>
    extract_docx.py --json --dry-run <input.docx> <out_dir>

출력:
    <out_dir>/text.txt        — 본문 텍스트 (단락 단위 줄바꿈)
    <out_dir>/img/<name>.png  — 임베디드 이미지 (또는 .jpg 등 원본 확장자)

종료 코드:
    0  성공 (또는 dry-run 검증 통과)
    1  입력 없음 / 인자 누락
    2  사용법 (--help 제외)

python-docx 가 없으면 zip + <w:t> 정규식 fallback. 그건 실패가 아니다
(DEP_MISS_PYTHON_DOCX 봉투, exit 0).
"""

from __future__ import annotations

import json
import os
import sys
import zipfile
from pathlib import Path


def emit(obj: dict, use_json: bool) -> None:
    if use_json:
        print(json.dumps(obj, ensure_ascii=False, indent=2))
    else:
        msg = obj.get("message") or obj.get("code")
        stream = sys.stdout if obj.get("ok") else sys.stderr
        print(msg, file=stream)


def main(argv: list[str] | None = None) -> int:
    argv = list(sys.argv[1:] if argv is None else argv)
    use_json = False
    dry = False
    while argv and argv[0].startswith("-"):
        flag = argv.pop(0)
        if flag == "--json":
            use_json = True
        elif flag == "--dry-run":
            dry = True
        elif flag in ("-h", "--help"):
            print(
                "사용법: extract_docx.py [--json] [--dry-run] <input.docx> <out_dir>",
                file=sys.stderr,
            )
            return 0
        else:
            emit(
                {
                    "schemaVersion": "1.0",
                    "helper": "extract_docx.py",
                    "ok": False,
                    "code": "DOCX_ARGS",
                    "message": f"오류: 알 수 없는 플래그 {flag}",
                },
                use_json,
            )
            return 2

    if len(argv) < 2:
        emit(
            {
                "schemaVersion": "1.0",
                "helper": "extract_docx.py",
                "ok": False,
                "code": "DOCX_ARGS",
                "message": "사용법: extract_docx.py [--json] [--dry-run] <input.docx> <out_dir>",
            },
            use_json,
        )
        return 1

    inp = Path(argv[0])
    out = Path(argv[1])

    if not inp.exists():
        emit(
            {
                "schemaVersion": "1.0",
                "helper": "extract_docx.py",
                "ok": False,
                "code": "DOCX_SRC_MISSING",
                "message": f"오류: 입력 파일이 없습니다: {inp}",
                "input": str(inp),
            },
            use_json,
        )
        return 1

    have_docx = False
    try:
        import docx as _docx  # type: ignore  # noqa: F401

        have_docx = True
    except ImportError:
        have_docx = False

    engine = "python-docx" if have_docx else "zip-regex-fallback"

    if dry:
        emit(
            {
                "schemaVersion": "1.0",
                "helper": "extract_docx.py",
                "ok": True,
                "code": "DOCX_OK",
                "dryRun": True,
                "engine": engine,
                "input": str(inp),
                "outDir": str(out),
                "planned": [
                    f"{out}/text.txt",
                    f"{out}/img/*",
                ],
                "pythonDocx": have_docx,
                "fallback": None if have_docx else "zip+<w:t> regex",
            },
            use_json,
        )
        return 0

    out.mkdir(parents=True, exist_ok=True)
    img_dir = out / "img"
    img_dir.mkdir(exist_ok=True)

    paragraph_count = 0
    # python-docx 우선 시도 (정밀 텍스트 추출)
    if have_docx:
        from docx import Document  # type: ignore

        d = Document(str(inp))
        with open(out / "text.txt", "w", encoding="utf-8") as f:
            for para in d.paragraphs:
                f.write(para.text + "\n")
        paragraph_count = len(d.paragraphs)
        if not use_json:
            print(f"텍스트 추출 ({paragraph_count} 단락): {out}/text.txt")
    else:
        # python-docx 없으면 zip 직접 파싱 (lxml 없이 정규식으로 단순 추출)
        import re

        with zipfile.ZipFile(inp, "r") as z:
            with z.open("word/document.xml") as f:
                xml = f.read().decode("utf-8", errors="ignore")
        texts = re.findall(r"<w:t[^>]*>([^<]*)</w:t>", xml)
        with open(out / "text.txt", "w", encoding="utf-8") as f:
            f.write("\n".join(texts))
        paragraph_count = len(texts)
        if not use_json:
            print(f"텍스트 추출 (정규식 fallback, {paragraph_count} 토큰): {out}/text.txt")

    # 임베디드 이미지 추출 (DOCX는 ZIP이므로 word/media/ 폴더에 들어있음)
    img_count = 0
    img_names: list[str] = []
    with zipfile.ZipFile(inp, "r") as z:
        for name in z.namelist():
            if name.startswith("word/media/"):
                base = os.path.basename(name)
                if not base:
                    continue
                with z.open(name) as src, open(img_dir / base, "wb") as dst:
                    dst.write(src.read())
                img_count += 1
                img_names.append(base)

    if use_json:
        emit(
            {
                "schemaVersion": "1.0",
                "helper": "extract_docx.py",
                "ok": True,
                "code": "DOCX_OK",
                "engine": engine,
                "paragraphs": paragraph_count,
                "images": img_count,
                "imageNames": img_names,
                "text": f"{out}/text.txt",
                "imgDir": str(img_dir),
                "pythonDocx": have_docx,
            },
            True,
        )
    else:
        print(f"이미지 추출 ({img_count} 개): {img_dir}/")
    return 0


if __name__ == "__main__":
    sys.exit(main())
